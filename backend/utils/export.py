"""Export utilities for generating PDF and Word documents from OCR text"""
import io
import os
import re
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND_NAME = "VERARA"
BRAND_TAGLINE = "POLYCLINIC & DIAGNOSTICS"


def _frontend_logo_path():
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    meipass = getattr(__import__("sys"), "_MEIPASS", "")
    candidates = [
        # Prefer rapha_logo.png (offline-safe, no internet needed)
        os.path.join(root, "assets", "rapha_logo.png"),
        os.path.join(root, "frontend", "public", "rapha_logo.png"),
        os.path.join(root, "frontend", "dist", "rapha_logo.png"),
        os.path.join(meipass, "rapha_logo.png") if meipass else "",
        # Fallback to existing jpg files
        os.path.join(root, "assets", "rapha_logo.jpg"),
        os.path.join(root, "frontend", "public", "rapha_logo.jpg"),
        os.path.join(root, "frontend", "dist", "rapha_logo.jpg"),
        os.path.join(root, "assets", "rapha_print_header.jpg"),
        os.path.join(root, "frontend", "public", "rapha_print_header.jpg"),
        os.path.join(root, "frontend", "dist", "rapha_print_header.jpg"),
        os.path.join(root, "frontend", "public", "logo.png"),
        os.path.join(root, "frontend", "dist", "logo.png"),
    ]
    return next((path for path in candidates if path and os.path.exists(path)), None)


def draw_pdf_brand_header(pdf, left_margin: float, top_y: float, title_width: float = 120):
    """Draw Rapha Medical Services branding header.

    Layout (left to right):
      - Logo image (22-24 mm wide) at x=left_margin, y=top_y
      - Hospital details text starting x=left_margin+26, y=top_y
      - Report title section is rendered separately by the caller (right-aligned)

    Never draws hospital text on top of separator lines or table borders.
    Caller must start first table at y >= 48 mm.
    """
    logo_path = _frontend_logo_path()
    logo_w = 22  # mm
    text_x = left_margin + logo_w + 4  # 4 mm gap after logo

    if logo_path:
        try:
            pdf.image(logo_path, x=left_margin, y=top_y, w=logo_w)
        except Exception:
            logo_path = None

    # Hospital details text (always rendered, with or without logo)
    if not logo_path:
        text_x = left_margin  # fallback: start at left margin

    pdf.set_xy(text_x, top_y + 0.5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(20, 61, 94)
    pdf.cell(70, 5.5, BRAND_NAME, ln=True)
    pdf.set_x(text_x)
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(20, 61, 94)
    pdf.cell(70, 4.5, "POLYCLINIC &", ln=True)
    pdf.set_x(text_x)
    pdf.cell(70, 4.5, "DIAGNOSTICS", ln=True)
    pdf.set_x(text_x)
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(70, 4, "Polyclinic | Diagnostics", ln=True)

    pdf.set_text_color(0, 0, 0)


def parse_markdown_line(line):
    """Parse a line and return (type, content, is_bold_segments)
    Types: header, bullet, table_row, text
    """
    line = line.strip()
    if not line:
        return ("empty", "", [])

    header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if header_match:
        return ("header", header_match.group(2).replace("**", ""), [])

    bullet_match = re.match(r"^[\•\-\*]\s+(.+)$", line)
    if bullet_match:
        content = bullet_match.group(1)
        return ("bullet", content, parse_bold_segments(content))

    if line.startswith("|") and line.endswith("|"):
        if re.match(r"^\|[\s\-\|:]+\|$", line):
            return ("table_sep", "", [])
        cells = [c.strip().replace("**", "") for c in line.strip("|").split("|")]
        return ("table_row", cells, [])

    return ("text", line, parse_bold_segments(line))


def parse_bold_segments(text):
    """Parse text and return list of (text, is_bold) tuples"""
    segments = []
    pattern = r"\*\*([^*]+)\*\*"
    last_end = 0

    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False))
        segments.append((match.group(1), True))
        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], False))

    if not segments:
        segments = [(text, False)]

    return segments


def generate_pdf(patient_name: str, doc_type: str, ocr_text: str, date_str: str = None) -> bytes:
    """Generate a hospital-standard A4 PDF with client hospital branding."""
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    class HospitalDocumentPDF(FPDF):
        def footer(self):
            self.set_y(-16)
            self.set_draw_color(215, 226, 236)
            self.line(15, self.get_y(), 195, self.get_y())
            self.set_y(-13)
            self.set_font("Helvetica", "", 7.5)
            self.set_text_color(90, 104, 119)
            self.cell(0, 4, _pdf_clean("This document is system generated for VERARA Polyclinic & Diagnostics. Verify clinical content before treatment decisions."), align="C", ln=True)
            self.cell(0, 4, _pdf_clean(f"Page {self.page_no()}/{{nb}}"), align="C")
            self.set_text_color(0, 0, 0)

    def _title_for_doc(value: str) -> str:
        normalized = (value or "").replace("-", "_").lower()
        if "prescription" in normalized:
            return "Doctor Prescription"
        if "discharge" in normalized:
            return "Discharge Summary"
        if "patient_journey" in normalized or "journey" in normalized:
            return "Patient Journey Report"
        return (value or "Medical Document").replace("_", " ").title()

    def _pdf_clean(value) -> str:
        text = "" if value is None else str(value)
        text = text.replace("₹", "Rs.").replace("•", "-").replace("–", "-").replace("—", "-")
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def _ensure_space(height: float):
        if pdf.get_y() + height > 276:
            pdf.add_page()
            _draw_document_header()

    def _cell_label_value(label: str, value: str, x: float, y: float, w: float, h: float = 12):
        pdf.set_xy(x, y)
        pdf.set_fill_color(248, 252, 255)
        pdf.set_draw_color(210, 226, 236)
        pdf.rect(x, y, w, h, style="DF")
        pdf.set_xy(x + 2.5, y + 2)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(59, 89, 111)
        pdf.cell(w - 5, 3.5, _pdf_clean(label.upper()), ln=True)
        pdf.set_x(x + 2.5)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(15, 23, 42)
        pdf.multi_cell(w - 5, 4, _pdf_clean(value or "-"))
        pdf.set_text_color(0, 0, 0)

    def _draw_document_header():
        pdf.set_y(13)
        draw_pdf_brand_header(pdf, left_margin, 13, title_width=92)
        pdf.set_xy(105, 15)
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(15, 45, 76)
        pdf.cell(90, 6, _pdf_clean(document_title), align="R", ln=True)
        pdf.set_x(105)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(83, 101, 120)
        pdf.cell(90, 5, _pdf_clean(f"Generated: {date_str}"), align="R", ln=True)
        pdf.set_x(105)
        pdf.cell(90, 5, _pdf_clean("VERARA clinical document"), align="R", ln=True)
        pdf.set_draw_color(13, 148, 169)
        pdf.set_line_width(0.6)
        pdf.line(left_margin, 42, 195, 42)
        pdf.set_text_color(0, 0, 0)
        pdf.set_y(48)

    def _section(title: str):
        _ensure_space(15)
        pdf.ln(2)
        pdf.set_fill_color(231, 247, 250)
        pdf.set_draw_color(190, 222, 233)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(11, 95, 118)
        pdf.cell(effective_width, 8, _pdf_clean(title), border=1, fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    def _paragraph(text: str):
        for raw in (text or "-").splitlines() or ["-"]:
            _ensure_space(8)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(effective_width, 5, _pdf_clean(raw.strip() or " "))
        pdf.set_text_color(0, 0, 0)

    def _two_column_details(rows):
        y = pdf.get_y()
        col_w = (effective_width - 4) / 2
        for index, (label, value) in enumerate(rows):
            if index % 2 == 0:
                _ensure_space(14)
                y = pdf.get_y()
            x = left_margin + (col_w + 4) * (index % 2)
            _cell_label_value(label, value, x, y, col_w, 13)
            if index % 2 == 1:
                pdf.set_y(y + 15)
        if len(rows) % 2 == 1:
            pdf.set_y(y + 15)

    def _markdown_table(cells):
        col_count = max(len(cells), 1)
        col_w = effective_width / col_count
        _ensure_space(8)
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(241, 248, 252)
        pdf.set_draw_color(215, 226, 236)
        for cell in cells:
            pdf.cell(col_w, 7, _pdf_clean(cell)[:24], border=1, fill=True)
        pdf.ln(7)

    pdf = HospitalDocumentPDF(format="A4")
    left_margin = 15
    right_margin = 15
    document_title = _title_for_doc(doc_type)
    pdf.set_left_margin(left_margin)
    pdf.set_right_margin(right_margin)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()
    pdf.add_page()

    page_width = pdf.w
    effective_width = page_width - left_margin - right_margin

    _draw_document_header()
    _section("Patient Information")
    _two_column_details([
        ("Patient Name", patient_name or "-"),
        ("Document Type", document_title),
        ("Generated On", date_str),
        ("Hospital System", BRAND_NAME),
    ])
    _section("Clinical Document Details")

    lines = ocr_text.split("\n")
    current_table_header = None

    for line in lines:
        line_type, content, segments = parse_markdown_line(line)

        if line_type == "empty":
            pdf.ln(2)
            continue

        if line_type == "header":
            current_table_header = None
            _section(content)
            continue

        if line_type == "bullet":
            _ensure_space(8)
            pdf.set_x(left_margin + 3)
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(5, 5, "-")
            for text, is_bold in segments:
                text = text.replace("**", "")
                pdf.set_font("Helvetica", "B" if is_bold else "", 10)
                pdf.write(5, _pdf_clean(text))
            pdf.ln(5)
            continue

        if line_type == "table_row":
            if current_table_header is None:
                current_table_header = content
                _markdown_table(content)
            else:
                col_count = max(len(current_table_header), len(content), 1)
                col_w = effective_width / col_count
                row_h = 7
                _ensure_space(row_h)
                pdf.set_font("Helvetica", "", 8)
                pdf.set_draw_color(225, 233, 240)
                for index in range(col_count):
                    value = content[index] if index < len(content) else ""
                    pdf.cell(col_w, row_h, _pdf_clean(value)[:28], border=1)
                pdf.ln(row_h)
            continue

        if line_type == "table_sep":
            continue

        if line_type == "text":
            current_table_header = None
            _ensure_space(8)
            for text, is_bold in segments:
                text = text.replace("**", "")
                pdf.set_font("Helvetica", "B" if is_bold else "", 10)
                pdf.write(5.3, _pdf_clean(text))
            pdf.ln(5.5)

    if "prescription" in (doc_type or "").lower():
        _section("Doctor Signature")
        pdf.ln(14)
        pdf.set_draw_color(120, 140, 160)
        pdf.line(132, pdf.get_y(), 190, pdf.get_y())
        pdf.set_xy(132, pdf.get_y() + 2)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(58, 5, _pdf_clean("Doctor Signature / Seal"), align="C")
    elif "discharge" in (doc_type or "").lower():
        _section("Doctor Certification")
        _paragraph("The above discharge summary has been prepared from available hospital records and clinical notes.")
        pdf.ln(10)
        pdf.set_draw_color(120, 140, 160)
        pdf.line(132, pdf.get_y(), 190, pdf.get_y())
        pdf.set_xy(132, pdf.get_y() + 2)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(58, 5, _pdf_clean("Consultant Signature / Seal"), align="C")

    return bytes(pdf.output())


def generate_word(patient_name: str, doc_type: str, ocr_text: str, date_str: str = None) -> bytes:
    """Generate Word document with OCR extracted text - properly formatted"""
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    doc = Document()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run(BRAND_NAME)
    brand_run.bold = True
    brand_run.font.size = Pt(24)
    brand_run.font.color.rgb = RGBColor(20, 61, 94)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(BRAND_TAGLINE)
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    title = doc.add_heading("Medical Document Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    cells = table.rows[0].cells
    cells[0].text = "Patient Name"
    cells[1].text = patient_name

    cells = table.rows[1].cells
    cells[0].text = "Document Type"
    cells[1].text = doc_type.replace("_", " ").title()

    cells = table.rows[2].cells
    cells[0].text = "Date"
    cells[1].text = date_str

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    lines = ocr_text.split("\n")

    for line in lines:
        line_type, content, segments = parse_markdown_line(line)

        if line_type == "empty":
            doc.add_paragraph()
            continue

        if line_type == "header":
            doc.add_heading(content, level=2)
            continue

        if line_type == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            for text, is_bold in segments:
                text = text.replace("**", "")
                run = para.add_run(text)
                run.bold = is_bold
            continue

        if line_type == "table_row":
            para = doc.add_paragraph()
            para.add_run("    ".join(content))
            continue

        if line_type == "table_sep":
            continue

        if line_type == "text":
            para = doc.add_paragraph()
            for text, is_bold in segments:
                text = text.replace("**", "")
                run = para.add_run(text)
                run.bold = is_bold

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
